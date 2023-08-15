import sqlite3 as sql
import os
import ctypes as ct
import docx
import datetime

parent_dir = os.path.expanduser('~')
_dir = 'billing_db'
_dir2 = 'billing_db\Bills'
path = os.path.join(parent_dir, _dir)
path2 = os.path.join(parent_dir, _dir2)


def initialize():
    try:
        mode = 0o666
        os.mkdir(mode=mode, path=path)
        os.mkdir(mode=mode, path=path2)
    except FileExistsError:
        pass
    global con
    global cur
    con = sql.connect(r'{}\z.db'.format(path))
    cur = con.cursor()
    con.execute('''CREATE TABLE IF NOT EXISTS Auth(username varchar(20) primary key not null, 
    password varchar(20) not null)''')
    con.execute("CREATE TABLE IF NOT EXISTS SignedIn(Boolean int not null)")
    con.execute('''CREATE TABLE IF NOT EXISTS Products(Product_ID integer primary key AUTOINCREMENT, 
    Product_Name varchar(30) not null, 
    Product_Type varchar(20), Cost int not null, Selling_Price int not null)''')
    con.execute('''CREATE TABLE IF NOT EXISTS Inventory(Product_ID int not null, product_name varchar(30) not null, 
    quantity int not null, FOREIGN KEY(Product_ID) REFERENCES Products(Product_ID) ON DELETE CASCADE)''')
    con.execute('''CREATE TABLE IF NOT EXISTS Profits(Order_Date DATE, Profit int)''')
    con.execute('''CREATE TABLE IF NOT EXISTS Customer(Customer_Name varchar(30), Profits int)''')
    con.commit()


def signed_in():
    cur.execute("SELECT * FROM SignedIn")
    signed_up = cur.fetchall()
    if len(signed_up) == 0:
        return False
    else:
        return True


def sign_up(username, password):
    con.execute('''CREATE TABLE IF NOT EXISTS Auth(Username varchar(20) not null, password varchar(20) not null)''')
    con.execute('''INSERT INTO Auth VALUES("{}", "{}")'''.format(username, password))
    con.execute('''INSERT INTO SignedIn VALUES(1)''')
    con.commit()


def login(username, password):
    cur.execute('''SELECT * FROM Auth WHERE username="{}" AND password="{}"'''.format(username, password))
    a = cur.fetchall()
    con.commit()
    if len(a) != 0:
        return True
    else:
        return False


def enter_products(values):
    for i in values:
        con.execute("INSERT INTO Products VALUES(NULL, '{}', '{}', {}, {})".format(i[0].title(),
                                                                                   i[1].title(), i[2], i[3]))
    con.commit()
    
    cur.execute("SELECT * FROM Products")
    a = len(cur.fetchall())
    b = len(values)
    cur.execute("SELECT Product_ID, Product_Name FROM Products where Product_ID>{}".format(a-b))
    c = cur.fetchall()
    for i in c:
        cur.execute("SELECT Product_ID FROM Inventory where Product_ID={}".format(i[0]))
        if len(cur.fetchall()) > 0:
            pass
        else:
            con.execute("INSERT INTO Inventory Values({}, '{}', 0)".format(i[0], i[1]))
    con.commit()


def check_products(values):
    for i in values:
        cur.execute("SELECT * FROM Products where Product_Name='{}'".format(i[0]))
        a = cur.fetchall()
        if len(a) > 0:
            return False
        else:
            continue
    return True


def update_inventory(values):
    for i in values:
        con.execute("update Inventory set quantity = quantity + {} where Product_ID={}".format(int(i[0]), int(i[1])))
    con.commit()


def check_data_integrity(values):
    for i in values:
        cur.execute('SELECT * FROM Products where Product_ID={}'.format(i[1]))
        a = cur.fetchall()
        if len(a) == 0:
            return False
        else:
            continue
    return True


def display_inventory():
    cur.execute('SELECT * FROM Inventory')
    return cur.fetchall()


def search(value):
    cur.execute("SELECT * FROM Inventory where product_name like '%{}%'".format(value.title()))
    return cur.fetchall()


def dark_title_bar(window):
    window.update()
    set_window_attribute = ct.windll.dwmapi.DwmSetWindowAttribute
    get_parent = ct.windll.user32.GetParent
    hwnd = get_parent(window.winfo_id())
    value = 2
    value = ct.c_int(value)
    set_window_attribute(hwnd, 20, ct.byref(value),
                         4)


def create_bill(data, customer_name="Customer", discount=0):
    date = str(datetime.date.today())
    date_time = datetime.datetime.now()
    time = date_time.strftime("%H;%M;%S")
    doc = docx.Document()
    doc.add_heading("XYZ Store", 0)
    doc.add_heading('''XYZ Street, New Delhi, Delhi''', 9)
    table = doc.add_table(rows=1, cols=5)
    row = table.rows[0].cells
    row[0].text = "Product ID"
    row[1].text = 'Product Name'
    row[2].text = 'Quantity'
    row[3].text = 'Price'
    row[4].text = "Net Cost"
    s = 0
    for Product_ID, Quant, Product, Price in data:
        row = table.add_row().cells
        row[0].text = str(Product_ID)
        row[1].text = str(Product)
        row[2].text = str(Quant)
        row[3].text = str(Price)
        row[4].text = str(Quant*Price)
        s += Quant*Price
    row2 = table.add_row().cells
    for i in range(4):
        row2[i].text = ''
    row2[3].text = 'Net+18% GST - Discount ='
    row2[4].text = str((1.18*s)*((100-discount)/100))
    table.style = "Light Shading Accent 1"
    para = doc.add_paragraph('''Thank you for shopping with XYZ Store! We hope you had a pleasant experience.
    We hope to see you again, ''')
    bold = para.add_run("Mr./Mrs. {}".format(customer_name))
    bold.bold = True
    doc.save(r"{}\{}({}, {}).docx".format(path2, customer_name, date, time))
    os.startfile(r"{}\{}({}, {}).docx".format(path2, customer_name, date, time), "open")


def fetch_data(value):
    cur.execute("SELECT Product_Name, Selling_Price, Cost from Products where Product_ID={}".format(value))
    val = cur.fetchall()
    return val


def reduce_inventory(data):
    for i in data:
        con.execute("update Inventory set quantity = quantity - {} where Product_ID={}".format(i[1], i[0]))
    con.commit()


def customer(value):
    con.execute("INSERT INTO Customer VALUES('{}', 0)".format(value))
    con.commit()


def customer_profit(value):
    con.execute("update Customer set Profits = Profits + {}".format(value))
    con.commit()


def customer_exists(value):
    cur.execute("SELECT * FROM Customer where Customer_Name='{}'".format(value))
    if len(cur.fetchall()) > 0:
        return True
    else:
        return False


def date_profits(values):
    con.execute("INSERT INTO Profits VALUES('{}', {})".format(values[0], values[1]))
    con.commit()
